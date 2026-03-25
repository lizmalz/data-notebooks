from docx import Document
from docx.enum.text import WD_BREAK

doc = Document("input.docx")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "Comment" in cell.text.lower():
                # Add a break after the row
                para = row.cells[-1].paragraphs[-1]
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)
                break  # move to next row once "Comment" is found

doc.save("output.docx")